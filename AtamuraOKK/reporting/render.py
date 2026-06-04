"""Render aggregated data + narrative into Markdown and .docx reports."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document

from AtamuraOKK.settings import settings

if TYPE_CHECKING:
    from AtamuraOKK.reporting.aggregate import ManagerReport, ReportData
    from AtamuraOKK.reporting.writer import ReportNarrative

_ZONE_TITLES = {
    "strong": "Сильная зона (85+)",
    "normal": "Нормальный уровень (80–84)",
    "borderline": "Пограничная зона (75–79)",
    "risk": "Зона риска (<75)",
}
_ZONE_ORDER = ["strong", "normal", "borderline", "risk"]


def _group_by_zone(managers: list[ManagerReport]) -> dict[str, list[ManagerReport]]:
    out: dict[str, list[ManagerReport]] = {z: [] for z in _ZONE_ORDER}
    for m in managers:
        out[m.zone].append(m)
    return out


def _zone_summary(data: ReportData) -> str:
    parts = [
        f"{_ZONE_TITLES[z].split(' (')[0]} — {data.zones.get(z, 0)}"
        for z in _ZONE_ORDER
    ]
    return ", ".join(parts)


def render_markdown(data: ReportData, narrative: ReportNarrative) -> str:  # noqa: C901
    """Render the report as Markdown."""
    notes = {n.name: n for n in narrative.manager_notes}
    norm = settings.report_score_norm
    targets = ", ".join(f"{k}: {v}" for k, v in data.targets.items())
    excluded = ", ".join(f"{k}: {v}" for k, v in data.excluded_by_type.items()) or "—"
    out: list[str] = [
        f"# Отчёт ОКК по телемаркетингу — {data.date_label}",
        "",
        "## Общая статистика",
        f"- Квалификационных звонков оценено: **{data.n_scored}**",
        f"- Средний балл команды: **{data.avg_percent}%** (норма {norm}%+)",
        "- Зоны: " + _zone_summary(data),
        "- Целевые/нецелевые: " + targets,
        f"- Звонков во флагах: **{data.n_flagged}**",
        f"- Исключено не по теме (напоминания/вендоры/внутренние/недозвоны): "
        f"**{data.n_excluded}** ({excluded})",
        "",
        "## Общая оценка",
        narrative.overall_assessment,
        "",
        "## Рейтинг менеджеров",
    ]
    for zone in _ZONE_ORDER:
        managers = _group_by_zone(data.managers)[zone]
        if not managers:
            continue
        out.append(f"\n### {_ZONE_TITLES[zone]}")
        for m in managers:
            out.append(f"\n**{m.name}** — {m.avg_percent}% ({m.n_calls} звонк.)")
            note = notes.get(m.name)
            if note:
                out.append(f"- Сильные стороны: {note.strengths}")
                out.append(f"- Зона роста: {note.growth_zone}")
                out.append(f"- Обучение: {note.training}")

    if narrative.systemic_problems:
        out.append("\n## Системные ошибки команды")
        out += [f"- {p}" for p in narrative.systemic_problems]

    if data.weakest_criteria:
        out.append("\n## Самые слабые критерии")
        out.append("\n| # | Блок | Критерий | Ср. % от макс |")
        out.append("|---|---|---|---|")
        for cr in data.weakest_criteria:
            out.append(
                f"| {cr.criterion_id} | {cr.block_name} | "
                f"{cr.criterion_text} | {cr.avg_pct_of_max}% |",
            )

    if data.flagged:
        out.append("\n## Проблемные звонки (флаги)")
        for c in data.flagged:
            flags = ", ".join(c.red_flags) or "—"
            out.append(
                f"- call #{c.call_id}: {c.percent}% [{c.zone}], {c.target_status}; "
                f"флаги: {flags}",
            )

    if narrative.recommendations:
        out.append("\n## Задачи и рекомендации для РОПов")
        out += [f"- {r}" for r in narrative.recommendations]

    out.append("\n## Итог")
    out.append(narrative.conclusion)
    return "\n".join(out) + "\n"


def render_docx(  # noqa: C901
    data: ReportData,
    narrative: ReportNarrative,
    path: Path,
) -> Path:
    """Render the report as a .docx file at ``path``."""
    notes = {n.name: n for n in narrative.manager_notes}
    doc = Document()
    doc.add_heading(f"Отчёт ОКК по телемаркетингу — {data.date_label}", level=0)

    doc.add_heading("Общая статистика", level=1)
    norm = settings.report_score_norm
    targets = ", ".join(f"{k}: {v}" for k, v in data.targets.items())
    excluded = ", ".join(f"{k}: {v}" for k, v in data.excluded_by_type.items()) or "—"
    for line in (
        f"Квалификационных звонков оценено: {data.n_scored}",
        f"Средний балл команды: {data.avg_percent}% (норма {norm}%+)",
        "Зоны: " + _zone_summary(data),
        "Целевые/нецелевые: " + targets,
        f"Звонков во флагах: {data.n_flagged}",
        f"Исключено не по теме: {data.n_excluded} ({excluded})",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Общая оценка", level=1)
    doc.add_paragraph(narrative.overall_assessment)

    doc.add_heading("Рейтинг менеджеров", level=1)
    grouped = _group_by_zone(data.managers)
    for zone in _ZONE_ORDER:
        if not grouped[zone]:
            continue
        doc.add_heading(_ZONE_TITLES[zone], level=2)
        for m in grouped[zone]:
            doc.add_paragraph().add_run(
                f"{m.name} — {m.avg_percent}% ({m.n_calls} звонк.)",
            ).bold = True
            note = notes.get(m.name)
            if note:
                doc.add_paragraph(
                    f"Сильные стороны: {note.strengths}", style="List Bullet"
                )
                doc.add_paragraph(
                    f"Зона роста: {note.growth_zone}", style="List Bullet"
                )
                doc.add_paragraph(f"Обучение: {note.training}", style="List Bullet")

    if narrative.systemic_problems:
        doc.add_heading("Системные ошибки команды", level=1)
        for p in narrative.systemic_problems:
            doc.add_paragraph(p, style="List Bullet")

    if data.weakest_criteria:
        doc.add_heading("Самые слабые критерии", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "#",
            "Блок",
            "Критерий",
            "Ср. % от макс",
        )
        for cr in data.weakest_criteria:
            cells = table.add_row().cells
            cells[0].text = str(cr.criterion_id)
            cells[1].text = cr.block_name
            cells[2].text = cr.criterion_text
            cells[3].text = f"{cr.avg_pct_of_max}%"

    if narrative.recommendations:
        doc.add_heading("Задачи и рекомендации для РОПов", level=1)
        for r in narrative.recommendations:
            doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Итог", level=1)
    doc.add_paragraph(narrative.conclusion)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
